import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_formatted_text(paragraph, text):
    # Split by bold markers (**text**)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
            tcBorders.append(element)

def parse_markdown_to_docx(md_path, docx_path):
    doc = Document()
    
    # Style configuration
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    table_buffer = []
    
    for line in lines:
        line = line.strip()
        
        # Table Processing
        if line.startswith('|'):
            table_buffer.append(line)
            continue
        else:
            if table_buffer:
                # Process the buffered table
                process_table(doc, table_buffer)
                table_buffer = []

        if not line:
            continue

        # Headers
        if line.startswith('#'):
            level = len(line.split(' ')[0])
            text = line.lstrip('#').strip()
            # Map MD levels to Docx levels (Title is usually larger, but let's map 1-1)
            # MD H1 -> Title often looks better, but let's stick to logical Headers
            if level == 1:
                p = doc.add_heading(text, 0) # Title style
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(text, level)
            continue

        # Blockquotes / Alerts
        if line.startswith('>'):
            text = line.lstrip('>').strip()
            # Detect alert types
            if text.startswith('[!') and text.endswith(']'):
                # It's an alert header, maybe bold it or ignore
                pass 
            else:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.italic = True
                p.style = 'Quote'
            continue
            
        # Lists
        if line.startswith('* ') or line.startswith('- '):
            text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            continue
            
        # Horizontal Rule
        if line.startswith('---'):
            doc.add_paragraph("_" * 70, style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        # Normal Text
        p = doc.add_paragraph()
        add_formatted_text(p, line)

    # Flush remaining table if any
    if table_buffer:
        process_table(doc, table_buffer)

    doc.save(docx_path)
    print(f"Document saved to {docx_path}")

def process_table(doc, table_lines):
    # Extract data
    rows = []
    for line in table_lines:
        # Split by | and remove empty start/end
        parts = [p.strip() for p in line.split('|') if p]
        rows.append(parts)
    
    if len(rows) < 2:
        return # Not a valid table

    # Remove separator line (e.g. ---|---|---)
    final_rows = [row for row in rows if not all('-' in cell for cell in row)]
    
    if not final_rows:
        return

    # Create table
    num_cols = len(final_rows[0])
    table = doc.add_table(rows=len(final_rows), cols=num_cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(final_rows):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            if j < len(row_cells):
                # Bold header row
                if i == 0:
                    p = row_cells[j].paragraphs[0]
                    run = p.add_run(cell_data.replace('**', '')) # unique bold handling for headers
                    run.bold = True
                else:
                    p = row_cells[j].paragraphs[0]
                    add_formatted_text(p, cell_data)

if __name__ == "__main__":
    md_file = 'h:\\works\\python\\Multi-Tenant\\EduERP_by_AI\\price.md'
    docx_file = 'h:\\works\\python\\Multi-Tenant\\EduERP_by_AI\\Codefyn_EduERP_Quotation.docx'
    parse_markdown_to_docx(md_file, docx_file)
